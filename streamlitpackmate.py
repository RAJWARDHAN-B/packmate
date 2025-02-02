import streamlit as st
import requests
import google.generativeai as genai
from groq import Groq
import re
from datetime import datetime
from io import BytesIO
from docx import Document

# Set page configuration
st.set_page_config(page_title="Smart Packing Assistant", page_icon="🎒", layout="wide")

# Load API keys
GEMINI_API_KEY = "AIzaSyBDEnO1lXyhUd6NctHbRqESI6BMdk61a8E"
GROQ_API_KEY = "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"
OPENCAGE_API_KEY = "993e21d6cca746a2bbebd2f6e02a8316"
WEATHER_API_KEY = "96ae2a7fe449f44ad93813d140e40aa1"

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# Cache for storing location coordinates
location_cache = {}


def get_lat_lon_from_opencage(location):
    if location in location_cache:
        return location_cache[location]
    url = f"https://api.opencagedata.com/geocode/v1/json?q={location}&key={OPENCAGE_API_KEY}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if data["results"]:
            lat = data["results"][0]["geometry"]["lat"]
            lon = data["results"][0]["geometry"]["lng"]
            location_cache[location] = (lat, lon)
            return lat, lon
    return None, None


def get_weather(lat, lon, date):
    url = f"https://api.openweathermap.org/data/2.5/forecast?lat={lat}&lon={lon}&appid={WEATHER_API_KEY}&units=metric"
    response = requests.get(url).json()
    if "list" in response:
        forecast = response["list"][0]
        description = forecast["weather"][0]["description"].capitalize()
        temperature = forecast["main"]["temp"]
        return f"{description}, {temperature}°C"
    return "Weather data not available"


def create_docx(weather, packing_list):
    """Generates a DOCX file with proper formatting."""
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Smart Packing Assistant", level=1)
    doc.add_paragraph(f"Weather Forecast: {weather}\n")
    doc.add_heading("Packing List:", level=2)

    for item in packing_list:
        doc.add_paragraph(f"- {item}")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def extract_items_from_suggestions(suggestions):
    """Cleans up AI-generated lists for better formatting."""
    clean_list = []
    for line in suggestions.split("\n"):
        line = re.sub(r"[:•\-]+", "", line).strip()
        if line and not line.isdigit():  # Avoid isolated numbers
            clean_list.append(line)
    return clean_list


def get_packing_suggestions(location, activities, trip_type, trip_duration, people):
    # Primary model (Gemini)
    model = genai.GenerativeModel(model_name="gemini-2.0-flash-exp")
    chat_session = model.start_chat(history=[])

    people_info = "; ".join(
        [f"{p['name']} (Age: {p['age']}, Gender: {p['gender']}, Medical Needs: {p['medical_issues']})" for p in people]
    )
    # Default to "general activities" if activities are not provided
    activities = activities if activities else "general activities"

    # Create the prompt dynamically based on trip type
    prompt = f"Suggest a packing list for travelers going to {location} for {activities}. The trip type is {trip_type}. " \
             f"Include items based on the trip duration of {trip_duration} days and travelers' details: {people_info}."

    try:
        response = chat_session.send_message(prompt)
        return extract_items_from_suggestions(response.text)
    except Exception as e:
        st.warning("❌ Gemini failed, falling back to LLaMA...")
        # Fallback model (LLaMA via Groq)
        client = Groq(api_key=GROQ_API_KEY)
        chat_session = client.chat.completions.create

        try:
            completion = chat_session.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": prompt}],
                temperature=1,
                max_completion_tokens=200,
                top_p=1,
                stream=False
            )
            return extract_items_from_suggestions(completion.choices[0].message.get("content", ""))
        except Exception as e:
            st.error(f"❌ Error generating packing list from both models: {e}")
            return []


def generate_minimalist_list(location, activities, trip_type, trip_duration, people):
    """Generate a minimalist packing list based on the trip details."""
    model = genai.GenerativeModel(model_name="gemini-2.0-flash-exp")
    chat_session = model.start_chat(history=[])

    people_info = "; ".join(
        [f"{p['name']} (Age: {p['age']}, Gender: {p['gender']}, Medical Needs: {p['medical_issues']})" for p in people]
    )
    activities = activities if activities else "general activities"

    # Generate a minimalist packing list prompt
    prompt = f"Suggest a minimalist packing list for travelers going to {location} for {activities}. The trip type is {trip_type}. " \
             f"Include only essential items based on the trip duration of {trip_duration} days and travelers' details: {people_info}. " \
             f"Make it lightweight and only include what is necessary."

    try:
        response = chat_session.send_message(prompt)
        return extract_items_from_suggestions(response.text)
    except Exception as e:
        st.warning("❌ Gemini failed, falling back to LLaMA...")
        # Fallback model (LLaMA via Groq)
        client = Groq(api_key=GROQ_API_KEY)
        chat_session = client.chat.completions.create

        try:
            completion = chat_session.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": prompt}],
                temperature=1,
                max_completion_tokens=200,
                top_p=1,
                stream=False
            )
            return extract_items_from_suggestions(completion.choices[0].message.get("content", ""))
        except Exception as e:
            st.error(f"❌ Error generating minimalist packing list from both models: {e}")
            return []


def filter_excessive_items(packing_list, trip_duration):
    """Limits excessive clothing quantities based on trip duration."""
    if isinstance(trip_duration, int):  # Temporary trip
        max_clothes = min(trip_duration + 2, 10)  # Pack max 10 clothing items
    else:
        max_clothes = 30  # Permanent relocation can pack more

    filtered_list = []
    for item in packing_list:
        match = re.search(r"(\d+)", item)
        if match and "TShirts" in item:
            item = re.sub(r"\d+", str(max_clothes), item)  # Adjust quantity
        filtered_list.append(item)

    return filtered_list


# Streamlit UI
st.title("🎒 Personalized Packing List Generator")
st.markdown("### Plan your trip smarter with weather-based personalized packing lists!")

st.sidebar.header("🌍 Trip Details")
location = st.sidebar.text_input("Enter a location:")
start_date = st.sidebar.date_input("Start Date", datetime.today())

# Dropdown for trip type options with Permanent Relocation
trip_type = st.sidebar.selectbox("Select Trip Type", [
    "Business Trip", "Permanent Relocation", "Solo Trip",
    "Educational Trip", "Romantic Trip", "Adventure Trip"
])

# Conditionally display End Date only if "Permanent Relocation" is not selected
end_date_container = st.sidebar.empty()
if trip_type != "Permanent Relocation":
    end_date = end_date_container.date_input("End Date", datetime.today())
else:
    end_date = None  # No End Date for Permanent Relocation

# Make activities field optional
activities = st.sidebar.text_area("Enter your activities (comma-separated):", value="")

# Choose between Minimalist or Detailed List
list_type = st.sidebar.selectbox("Select Packing List Type", ["Detailed List", "Minimalist List"])

num_people = st.sidebar.number_input("Number of people:", min_value=1, step=1)

people = []
for i in range(num_people):
    with st.expander(f"👤 Person {i + 1} Details"):
        name = st.text_input(f"Name", key=f"name_{i}")
        gender = st.selectbox("Gender", ["Male", "Female", "Other"], key=f"gender_{i}")
        age = st.number_input("Age", min_value=0, step=1, key=f"age_{i}")
        medical_issues = st.text_area("Any medical issues", key=f"medical_{i}")
        people.append({"name": name, "gender": gender, "age": age, "medical_issues": medical_issues})

if st.sidebar.button("Generate Packing List 🧳"):
    if not location or not people:
        st.error("🚨 Please enter all required fields.")
    else:
        lat, lon = get_lat_lon_from_opencage(location)
        if lat and lon:
            if trip_type != "Permanent Relocation":
                trip_duration = max(1, (end_date - start_date).days)  # Ensure at least 1 day
            else:
                trip_duration = "Permanent"

            weather = get_weather(lat, lon, start_date)

            # Generate packing list based on selected list type
            if list_type == "Minimalist List":
                packing_list = generate_minimalist_list(location, activities, trip_type, trip_duration, people)
            else:
                packing_list = get_packing_suggestions(location, activities, trip_type, trip_duration, people)

            # Apply item filtering
            packing_list = filter_excessive_items(packing_list, trip_duration)

            st.success("✅ Packing list generated successfully!")
            st.subheader("🌤 Weather Forecast")
            st.info(weather)
            st.subheader(f"📦 {list_type}")
            for item in packing_list:
                st.write(f"- {item}")

            docx_file = create_docx(weather, packing_list)
            st.download_button("📥 Download Packing List", data=docx_file, file_name="Packing_List.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.error("❌ Unable to retrieve location coordinates.")
